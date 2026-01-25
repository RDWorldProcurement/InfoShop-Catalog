"""
Real Document Extraction Module
Extracts quotation data from uploaded files (PDF, images, Excel, Word) using AI
"""

import os
import io
import json
import base64
import logging
from typing import Dict, Optional, List
from datetime import datetime
from dotenv import load_dotenv

load_dotenv()

logger = logging.getLogger(__name__)

# Import file processing libraries
try:
    from PyPDF2 import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logger.warning("PyPDF2 not available - PDF text extraction disabled")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available - DOCX extraction disabled")

try:
    from openpyxl import load_workbook
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    logger.warning("openpyxl not available - Excel extraction disabled")

try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False
    logger.warning("Pillow not available - Image processing disabled")

# Import Emergent LLM integration
try:
    from emergentintegrations.llm.chat import LlmChat, UserMessage
    EMERGENT_AVAILABLE = True
except ImportError:
    EMERGENT_AVAILABLE = False
    logger.warning("emergentintegrations not available")

EMERGENT_LLM_KEY = os.environ.get("EMERGENT_LLM_KEY")

# System prompt for document extraction
EXTRACTION_SYSTEM_PROMPT = """You are an expert document parser specializing in extracting structured data from procurement quotations and invoices.

Your task is to extract ALL information from the provided document content and return it in a specific JSON format.

IMPORTANT RULES:
1. Extract EXACTLY what is in the document - do not invent or assume data
2. If a field is not present in the document, use null or appropriate default
3. Parse ALL line items with their exact descriptions, quantities, unit prices
4. Identify the supplier/vendor information
5. Extract totals, taxes, and any discounts mentioned
6. Note the currency if specified
7. For EACH line item, classify with UNSPSC code based on product/service description

You MUST respond with ONLY valid JSON in this exact structure:
{
    "supplier": {
        "name": "extracted supplier/vendor name",
        "address": "full address if available",
        "city": "city, state/country",
        "tax_id": "tax ID if present",
        "contact_email": "email if present",
        "contact_phone": "phone if present"
    },
    "quotation_details": {
        "quotation_number": "quote/invoice number",
        "quotation_date": "YYYY-MM-DD format",
        "valid_until": "YYYY-MM-DD if present",
        "payment_terms": "payment terms if specified",
        "delivery_terms": "delivery terms if specified",
        "currency": "USD/EUR/etc"
    },
    "line_items": [
        {
            "line_number": 1,
            "description": "exact product/service description from document",
            "quantity": number,
            "unit_price": number,
            "unit": "EA/HR/etc",
            "line_total": number,
            "category": "best category match",
            "part_number": "if present",
            "unspsc_code": "8-digit UNSPSC code (e.g., 31171500 for Bearings)",
            "unspsc_category": "UNSPSC category name"
        }
    ],
    "totals": {
        "subtotal": number,
        "tax_rate": decimal (e.g., 0.08 for 8%),
        "tax_amount": number,
        "shipping": number,
        "discount": number,
        "grand_total": number
    },
    "notes": "any additional notes or terms from the document",
    "extraction_confidence": 0.0-1.0 based on clarity of document
}

If the document is unclear or you cannot extract certain information, still provide the structure with null values and note the issues in the notes field."""


def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text content from PDF file"""
    if not PDF_AVAILABLE:
        return ""
    
    try:
        pdf_reader = PdfReader(io.BytesIO(file_content))
        text_content = []
        for page in pdf_reader.pages:
            text = page.extract_text()
            if text:
                text_content.append(text)
        return "\n\n".join(text_content)
    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        return ""


def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text content from Word document"""
    if not DOCX_AVAILABLE:
        return ""
    
    try:
        doc = Document(io.BytesIO(file_content))
        text_content = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text)
        
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_content.append(row_text)
        
        return "\n".join(text_content)
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        return ""


def extract_text_from_excel(file_content: bytes) -> str:
    """Extract text content from Excel file"""
    if not EXCEL_AVAILABLE:
        return ""
    
    try:
        workbook = load_workbook(io.BytesIO(file_content), data_only=True)
        text_content = []
        
        for sheet in workbook.worksheets:
            text_content.append(f"=== Sheet: {sheet.title} ===")
            for row in sheet.iter_rows():
                row_values = []
                for cell in row:
                    if cell.value is not None:
                        row_values.append(str(cell.value))
                if row_values:
                    text_content.append(" | ".join(row_values))
        
        return "\n".join(text_content)
    except Exception as e:
        logger.error(f"Excel extraction error: {e}")
        return ""


def image_to_base64(file_content: bytes, file_type: str) -> str:
    """Convert image to base64 for AI vision processing"""
    if not PIL_AVAILABLE:
        return ""
    
    try:
        # Resize large images to reduce token usage
        image = Image.open(io.BytesIO(file_content))
        
        # Convert to RGB if necessary
        if image.mode in ('RGBA', 'P'):
            image = image.convert('RGB')
        
        # Resize if too large (max 2000px on longest side)
        max_size = 2000
        if max(image.size) > max_size:
            ratio = max_size / max(image.size)
            new_size = tuple(int(dim * ratio) for dim in image.size)
            image = image.resize(new_size, Image.Resampling.LANCZOS)
        
        # Convert to base64
        buffer = io.BytesIO()
        image.save(buffer, format='JPEG', quality=85)
        return base64.b64encode(buffer.getvalue()).decode('utf-8')
    except Exception as e:
        logger.error(f"Image processing error: {e}")
        return ""


async def extract_with_ai_text(text_content: str, session_id: str) -> Dict:
    """Use AI to extract structured data from text content"""
    if not EMERGENT_AVAILABLE or not EMERGENT_LLM_KEY:
        logger.error("Emergent LLM not available for extraction")
        return None
    
    if not text_content or len(text_content.strip()) < 50:
        logger.warning("Insufficient text content for extraction")
        return None
    
    try:
        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=f"{session_id}_extract",
            system_message=EXTRACTION_SYSTEM_PROMPT
        ).with_model("openai", "gpt-5.2")
        
        prompt = f"""Please extract all quotation/invoice data from the following document content:

---DOCUMENT START---
{text_content[:15000]}  
---DOCUMENT END---

Extract all line items, supplier info, totals, and return as JSON."""

        message = UserMessage(text=prompt)
        response = await chat.send_message(message)
        
        # Parse JSON response
        response_text = str(response)
        
        # Try to extract JSON from response
        if "```json" in response_text:
            json_str = response_text.split("```json")[1].split("```")[0].strip()
        elif "```" in response_text:
            json_str = response_text.split("```")[1].split("```")[0].strip()
        else:
            # Find JSON object in response
            start_idx = response_text.find('{')
            end_idx = response_text.rfind('}') + 1
            if start_idx >= 0 and end_idx > start_idx:
                json_str = response_text[start_idx:end_idx]
            else:
                json_str = response_text
        
        extracted_data = json.loads(json_str)
        extracted_data["extraction_method"] = "ai_text"
        return extracted_data
        
    except json.JSONDecodeError as e:
        logger.error(f"JSON parsing error: {e}")
        return None
    except Exception as e:
        logger.error(f"AI text extraction error: {e}")
        return None


async def extract_with_ai_vision(image_base64: str, session_id: str) -> Dict:
    """Use AI vision to extract structured data from image"""
    if not EMERGENT_AVAILABLE or not EMERGENT_LLM_KEY:
        logger.error("Emergent LLM not available for extraction")
        return None
    
    try:
        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=f"{session_id}_vision",
            system_message=EXTRACTION_SYSTEM_PROMPT
        ).with_model("openai", "gpt-5.2")
        
        # Create message with image
        message = UserMessage(
            text="Please extract all quotation/invoice data from this document image. Return the data as JSON.",
            images=[f"data:image/jpeg;base64,{image_base64}"]
        )
        response = await chat.send_message(message)
        
        # Parse JSON response
        response_text = str(response)
        
        if "```json" in response_text:
            json_str = response_text.split("```json")[1].split("```")[0].strip()
        elif "```" in response_text:
            json_str = response_text.split("```")[1].split("```")[0].strip()
        else:
            start_idx = response_text.find('{')
            end_idx = response_text.rfind('}') + 1
            if start_idx >= 0 and end_idx > start_idx:
                json_str = response_text[start_idx:end_idx]
            else:
                json_str = response_text
        
        extracted_data = json.loads(json_str)
        extracted_data["extraction_method"] = "ai_vision"
        return extracted_data
        
    except json.JSONDecodeError as e:
        logger.error(f"JSON parsing error in vision: {e}")
        return None
    except Exception as e:
        logger.error(f"AI vision extraction error: {e}")
        return None


async def extract_quotation_data(
    file_content: bytes,
    file_name: str,
    file_type: str,
    supplier_name: Optional[str] = None,
    session_id: str = "extract"
) -> Dict:
    """
    Main function to extract quotation data from uploaded file.
    Supports PDF, images, Excel, and Word documents.
    """
    logger.info(f"Extracting data from {file_name} (type: {file_type})")
    
    extracted_data = None
    text_content = ""
    
    # Determine file type and extract accordingly
    file_ext = file_name.lower().split('.')[-1] if '.' in file_name else ''
    
    # Try text extraction first for supported formats
    if file_type == 'application/pdf' or file_ext == 'pdf':
        text_content = extract_text_from_pdf(file_content)
        logger.info(f"Extracted {len(text_content)} chars from PDF")
        
    elif file_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 
                       'application/msword'] or file_ext in ['docx', 'doc']:
        text_content = extract_text_from_docx(file_content)
        logger.info(f"Extracted {len(text_content)} chars from Word doc")
        
    elif file_type in ['application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                       'application/vnd.ms-excel'] or file_ext in ['xlsx', 'xls']:
        text_content = extract_text_from_excel(file_content)
        logger.info(f"Extracted {len(text_content)} chars from Excel")
    
    # If we have text content, use text-based AI extraction
    if text_content and len(text_content.strip()) > 100:
        logger.info("Using AI text extraction")
        extracted_data = await extract_with_ai_text(text_content, session_id)
    
    # For images or if text extraction failed, use vision
    if extracted_data is None:
        is_image = (file_type.startswith('image/') or 
                    file_ext in ['png', 'jpg', 'jpeg', 'gif', 'webp', 'tiff'])
        
        if is_image or (not text_content and PDF_AVAILABLE):
            logger.info("Using AI vision extraction")
            
            # For PDFs without text, convert to image
            if file_type == 'application/pdf' or file_ext == 'pdf':
                # Try text one more time, if still empty, we'd need pdf2image
                # For now, return error suggesting image upload
                if not text_content:
                    logger.warning("PDF has no extractable text - may be scanned document")
            
            # For actual images
            if is_image:
                image_b64 = image_to_base64(file_content, file_type)
                if image_b64:
                    extracted_data = await extract_with_ai_vision(image_b64, session_id)
    
    # If extraction failed, return error structure
    if extracted_data is None:
        logger.error("All extraction methods failed")
        return {
            "error": True,
            "message": "Could not extract data from document. Please ensure the file contains readable text or is a clear image.",
            "supplier": {
                "name": supplier_name or "Unknown",
                "address": None,
                "city": None,
                "tax_id": None,
                "contact_email": None
            },
            "quotation_details": {
                "quotation_number": None,
                "quotation_date": datetime.now().strftime("%Y-%m-%d"),
                "valid_until": None,
                "payment_terms": None,
                "delivery_terms": None
            },
            "line_items": [],
            "totals": {
                "subtotal": 0,
                "tax_rate": 0,
                "tax_amount": 0,
                "shipping": 0,
                "grand_total": 0
            },
            "extraction_confidence": 0,
            "document_language": "Unknown",
            "pages_processed": 0
        }
    
    # Validate and clean up extracted data
    extracted_data = validate_and_clean_extraction(extracted_data, supplier_name)
    
    return extracted_data


def validate_and_clean_extraction(data: Dict, supplier_name: Optional[str] = None) -> Dict:
    """Validate and clean up extracted data, ensuring proper structure"""
    
    # Ensure supplier structure
    if "supplier" not in data or not isinstance(data["supplier"], dict):
        data["supplier"] = {}
    
    supplier = data["supplier"]
    supplier.setdefault("name", supplier_name or "Unknown Supplier")
    supplier.setdefault("address", None)
    supplier.setdefault("city", None)
    supplier.setdefault("tax_id", None)
    supplier.setdefault("contact_email", None)
    
    # Override with provided supplier name if given
    if supplier_name:
        supplier["name"] = supplier_name
    
    # Ensure quotation_details structure
    if "quotation_details" not in data or not isinstance(data["quotation_details"], dict):
        data["quotation_details"] = {}
    
    details = data["quotation_details"]
    details.setdefault("quotation_number", f"QT-{datetime.now().strftime('%Y%m%d%H%M%S')}")
    details.setdefault("quotation_date", datetime.now().strftime("%Y-%m-%d"))
    details.setdefault("valid_until", None)
    details.setdefault("payment_terms", None)
    details.setdefault("delivery_terms", None)
    details.setdefault("currency", "USD")
    
    # Ensure line_items is a list with proper structure
    if "line_items" not in data or not isinstance(data["line_items"], list):
        data["line_items"] = []
    
    cleaned_items = []
    for i, item in enumerate(data["line_items"]):
        if not isinstance(item, dict):
            continue
        
        cleaned_item = {
            "line_number": item.get("line_number", i + 1),
            "description": str(item.get("description", "Unknown Item")),
            "quantity": float(item.get("quantity", 1)) if item.get("quantity") else 1,
            "unit_price": float(item.get("unit_price", 0)) if item.get("unit_price") else 0,
            "unit": str(item.get("unit", "EA")),
            "line_total": float(item.get("line_total", 0)) if item.get("line_total") else 0,
            "category": str(item.get("category", "General")),
            "part_number": item.get("part_number")
        }
        
        # Calculate line_total if not provided
        if cleaned_item["line_total"] == 0 and cleaned_item["unit_price"] > 0:
            cleaned_item["line_total"] = round(cleaned_item["unit_price"] * cleaned_item["quantity"], 2)
        
        cleaned_items.append(cleaned_item)
    
    data["line_items"] = cleaned_items
    
    # Ensure totals structure
    if "totals" not in data or not isinstance(data["totals"], dict):
        data["totals"] = {}
    
    totals = data["totals"]
    
    # Calculate subtotal from line items if not provided
    calculated_subtotal = sum(item["line_total"] for item in cleaned_items)
    totals.setdefault("subtotal", calculated_subtotal or 0)
    totals.setdefault("tax_rate", 0)
    totals.setdefault("tax_amount", 0)
    totals.setdefault("shipping", 0)
    totals.setdefault("discount", 0)
    
    # Calculate grand_total if not provided
    if "grand_total" not in totals or not totals["grand_total"]:
        totals["grand_total"] = (
            float(totals.get("subtotal", 0)) + 
            float(totals.get("tax_amount", 0)) + 
            float(totals.get("shipping", 0)) - 
            float(totals.get("discount", 0))
        )
    
    # Convert all numeric values to proper types
    for key in ["subtotal", "tax_rate", "tax_amount", "shipping", "discount", "grand_total"]:
        if key in totals:
            try:
                totals[key] = float(totals[key]) if totals[key] else 0
            except (ValueError, TypeError):
                totals[key] = 0
    
    # Add metadata
    data.setdefault("extraction_confidence", 0.85)
    data.setdefault("document_language", "English")
    data.setdefault("pages_processed", 1)
    data.setdefault("error", False)
    
    return data


# UNSPSC Code Reference for common procurement categories
UNSPSC_REFERENCE = {
    # MRO & Industrial
    "31170000": {"name": "Bearings and bushings and wheels and gears", "keywords": ["bearing", "bushing", "wheel", "gear", "roller", "ball bearing"]},
    "31171500": {"name": "Ball bearings", "keywords": ["ball bearing", "radial bearing", "angular bearing"]},
    "31171600": {"name": "Roller bearings", "keywords": ["roller bearing", "needle bearing", "tapered roller"]},
    "39110000": {"name": "Lamps and lightbulbs and lamp components", "keywords": ["lamp", "light", "bulb", "led", "fluorescent", "lighting"]},
    "31160000": {"name": "Hardware and fasteners", "keywords": ["fastener", "bolt", "screw", "nut", "washer", "anchor", "rivet"]},
    "27110000": {"name": "Hand tools", "keywords": ["wrench", "screwdriver", "plier", "hammer", "saw", "hand tool"]},
    "27112000": {"name": "Power tools", "keywords": ["drill", "grinder", "power tool", "electric tool", "cordless"]},
    "46180000": {"name": "Safety and protection", "keywords": ["safety", "ppe", "helmet", "glove", "goggle", "vest", "protection", "respirator"]},
    "31190000": {"name": "Abrasives and abrasive media", "keywords": ["abrasive", "sandpaper", "grinding wheel", "polishing"]},
    "31200000": {"name": "Adhesives and sealants", "keywords": ["adhesive", "sealant", "glue", "epoxy", "silicone"]},
    "47130000": {"name": "Cleaning equipment and supplies", "keywords": ["cleaning", "janitorial", "mop", "broom", "cleaner", "detergent"]},
    "40100000": {"name": "Heating and ventilation and air circulation", "keywords": ["hvac", "air conditioning", "ventilation", "heating", "fan", "blower"]},
    "40140000": {"name": "Fluid and gas distribution", "keywords": ["hydraulic", "pneumatic", "valve", "pump", "pipe", "fitting", "plumbing"]},
    "41110000": {"name": "Laboratory and measuring equipment", "keywords": ["laboratory", "lab", "test", "measurement", "meter", "gauge", "instrument"]},
    "15120000": {"name": "Lubricants and oils and greases", "keywords": ["lubricant", "oil", "grease", "lubrication"]},
    "24100000": {"name": "Material handling machinery", "keywords": ["forklift", "pallet", "conveyor", "hoist", "crane", "material handling"]},
    "26100000": {"name": "Power sources", "keywords": ["motor", "engine", "generator", "drive", "vfd", "inverter"]},
    "24110000": {"name": "Containers and packaging", "keywords": ["packaging", "container", "box", "crate", "pallet", "shipping"]},
    "40170000": {"name": "Pipe fittings", "keywords": ["pipe", "fitting", "coupling", "elbow", "tee", "flange"]},
    "40150000": {"name": "Industrial pumps and compressors", "keywords": ["pump", "compressor", "vacuum"]},
    "11100000": {"name": "Minerals and ores and metals", "keywords": ["steel", "aluminum", "metal", "alloy", "raw material"]},
    "23270000": {"name": "Welding and soldering equipment", "keywords": ["welding", "welder", "solder", "torch", "electrode"]},
    "32150000": {"name": "Automation and control equipment", "keywords": ["plc", "automation", "control", "sensor", "actuator", "relay"]},
    
    # IT Equipment
    "43211500": {"name": "Computers", "keywords": ["computer", "laptop", "notebook", "desktop", "pc", "workstation"]},
    "43211900": {"name": "Computer displays", "keywords": ["monitor", "display", "screen", "lcd", "led display"]},
    "43222600": {"name": "Network hardware", "keywords": ["router", "switch", "network", "ethernet", "wifi", "access point"]},
    "43211800": {"name": "Computer servers", "keywords": ["server", "rack server", "blade server", "storage"]},
    "43211700": {"name": "Computer accessories", "keywords": ["keyboard", "mouse", "webcam", "headset", "usb"]},
    "43233000": {"name": "Software", "keywords": ["software", "license", "subscription", "saas", "application"]},
    "43232600": {"name": "Storage devices", "keywords": ["hard drive", "ssd", "storage", "nas", "backup"]},
    
    # Professional Services
    "80100000": {"name": "Management advisory services", "keywords": ["consulting", "advisory", "management", "strategy"]},
    "80110000": {"name": "Human resources services", "keywords": ["hr", "recruitment", "staffing", "training", "human resource"]},
    "81110000": {"name": "Computer services", "keywords": ["it services", "it support", "managed services", "technical support"]},
    "81112000": {"name": "Software maintenance and support", "keywords": ["software support", "maintenance", "upgrade", "patch"]},
    "82100000": {"name": "Advertising", "keywords": ["advertising", "marketing", "campaign", "media buying"]},
    "83100000": {"name": "Utilities", "keywords": ["electricity", "gas", "water", "utility"]},
    "84110000": {"name": "Accounting services", "keywords": ["accounting", "audit", "bookkeeping", "tax", "financial"]},
    "85120000": {"name": "Healthcare services", "keywords": ["medical", "health", "clinical", "hospital"]},
    "86130000": {"name": "Training services", "keywords": ["training", "education", "course", "workshop", "certification"]},
    "90100000": {"name": "Restaurants and catering", "keywords": ["catering", "food service", "cafeteria", "meal"]},
    "72150000": {"name": "Production and manufacturing", "keywords": ["manufacturing", "production", "assembly", "fabrication"]},
    "78100000": {"name": "Mail and cargo transport", "keywords": ["shipping", "freight", "logistics", "transportation", "courier"]},
}


async def classify_unspsc_with_ai(line_items: List[Dict], session_id: str = None) -> List[Dict]:
    """
    Use AI to classify line items with UNSPSC codes.
    This performs deep semantic matching beyond simple keyword search.
    """
    if not line_items:
        return line_items
    
    if not EMERGENT_AVAILABLE or not EMERGENT_LLM_KEY:
        # Fallback to keyword-based classification
        return classify_unspsc_by_keywords(line_items)
    
    try:
        # Prepare items for classification
        items_text = "\n".join([
            f"{i+1}. {item.get('description', 'Unknown')} (Qty: {item.get('quantity', 1)}, Unit: {item.get('unit', 'EA')})"
            for i, item in enumerate(line_items)
        ])
        
        unspsc_system_prompt = """You are an expert UNSPSC (United Nations Standard Products and Services Code) classifier.
Your task is to assign the most accurate 8-digit UNSPSC code to each product or service.

UNSPSC Structure:
- Segment (2 digits): Broad category
- Family (4 digits): Subcategory
- Class (6 digits): Product group
- Commodity (8 digits): Specific product

Common UNSPSC Segments:
- 10-15: Raw materials
- 20-27: Industrial equipment and tools
- 30-32: Components and supplies
- 39-47: Facility maintenance and supplies
- 43: IT equipment and software
- 72-86: Services

Return ONLY valid JSON array with this structure for each item:
[
    {
        "item_index": 1,
        "description": "original description",
        "unspsc_code": "8-digit code",
        "unspsc_segment": "2-digit segment",
        "unspsc_family": "4-digit family",
        "unspsc_class": "6-digit class",
        "unspsc_category": "Full category name",
        "classification_confidence": 0.0-1.0,
        "classification_rationale": "Brief explanation"
    }
]"""

        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=session_id or f"unspsc_classify_{datetime.now().timestamp()}",
            system_message=unspsc_system_prompt
        ).with_model("openai", "gpt-5.2")
        
        prompt = f"""Classify the following {len(line_items)} items with their UNSPSC codes:

{items_text}

For each item, determine the most specific 8-digit UNSPSC code. Consider:
1. The exact product/service type
2. The industry context (industrial, IT, services)
3. Common procurement classifications

Return ONLY the JSON array."""

        response = await chat.send_message(UserMessage(text=prompt))
        response_text = str(response)
        
        # Parse JSON response
        try:
            if "```json" in response_text:
                json_str = response_text.split("```json")[1].split("```")[0].strip()
            elif "```" in response_text:
                json_str = response_text.split("```")[1].split("```")[0].strip()
            elif "[" in response_text:
                start = response_text.find("[")
                end = response_text.rfind("]") + 1
                json_str = response_text[start:end]
            else:
                raise ValueError("No JSON found")
            
            classifications = json.loads(json_str)
            
            # Merge classifications back into line items
            for classification in classifications:
                idx = classification.get("item_index", 1) - 1
                if 0 <= idx < len(line_items):
                    line_items[idx]["unspsc_code"] = classification.get("unspsc_code", "00000000")
                    line_items[idx]["unspsc_category"] = classification.get("unspsc_category", "Unclassified")
                    line_items[idx]["unspsc_segment"] = classification.get("unspsc_segment")
                    line_items[idx]["unspsc_family"] = classification.get("unspsc_family")
                    line_items[idx]["unspsc_class"] = classification.get("unspsc_class")
                    line_items[idx]["classification_confidence"] = classification.get("classification_confidence", 0.85)
            
            logger.info(f"AI UNSPSC classification completed for {len(line_items)} items")
            return line_items
            
        except json.JSONDecodeError as e:
            logger.warning(f"UNSPSC JSON parse error: {e}")
            return classify_unspsc_by_keywords(line_items)
            
    except Exception as e:
        logger.error(f"AI UNSPSC classification error: {e}")
        return classify_unspsc_by_keywords(line_items)


def classify_unspsc_by_keywords(line_items: List[Dict]) -> List[Dict]:
    """
    Fallback keyword-based UNSPSC classification.
    Matches item descriptions against known UNSPSC categories.
    """
    for item in line_items:
        description = item.get("description", "").lower()
        best_match = None
        best_score = 0
        
        for code, info in UNSPSC_REFERENCE.items():
            score = 0
            for keyword in info["keywords"]:
                if keyword.lower() in description:
                    score += len(keyword)  # Longer keyword matches score higher
            
            if score > best_score:
                best_score = score
                best_match = (code, info["name"])
        
        if best_match:
            item["unspsc_code"] = best_match[0]
            item["unspsc_category"] = best_match[1]
            item["classification_confidence"] = min(0.6 + (best_score * 0.05), 0.85)
        else:
            # Default to general supplies
            item["unspsc_code"] = "00000000"
            item["unspsc_category"] = "Unclassified - Manual Review Required"
            item["classification_confidence"] = 0.3
    
    return line_items
